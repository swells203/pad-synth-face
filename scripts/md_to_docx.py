"""One-off Markdown -> .docx converter for the DefinitiveID design doc.

Handles the subset of Markdown actually used in the spec: pandoc-style
title block (% lines), ATX headings, bold/italic inline, bullet and
ordered lists, fenced code blocks, pipe tables, and horizontal rules.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _add_inline(paragraph, text: str) -> None:
    """Render **bold** segments; everything else plain."""
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        run = paragraph.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _is_table_sep(line: str) -> bool:
    s = line.strip()
    return bool(s) and set(s) <= set("|:- ") and "-" in s


def convert(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text().splitlines()
    doc = Document()

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    n = len(lines)

    # Pandoc title block: leading lines starting with "% "
    title_block: list[str] = []
    while i < n and lines[i].startswith("% "):
        title_block.append(lines[i][2:].strip())
        i += 1
    if title_block:
        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run(title_block[0])
        r.bold = True
        r.font.size = Pt(22)
        for sub in title_block[1:]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = p.add_run(sub)
            rr.font.size = Pt(13)
            rr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        doc.add_page_break()

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Skip the YAML-ish "---" rules and blank lines
        if stripped == "---":
            i += 1
            continue
        if stripped == "":
            i += 1
            continue

        # Fenced code block
        if stripped.startswith("```"):
            i += 1
            code: list[str] = []
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1  # consume closing fence
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code))
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
            continue

        # Headings
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        # Pipe table
        if stripped.startswith("|") and i + 1 < n and _is_table_sep(lines[i + 1]):
            header = [c.strip() for c in stripped.strip("|").split("|")]
            i += 2  # skip header + separator
            rows: list[list[str]] = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for j, h in enumerate(header):
                cell = table.rows[0].cells[j]
                cell.text = ""
                _add_inline(cell.paragraphs[0], h)
                for r in cell.paragraphs[0].runs:
                    r.bold = True
            for row in rows:
                cells = table.add_row().cells
                for j in range(len(header)):
                    val = row[j] if j < len(row) else ""
                    cells[j].text = ""
                    _add_inline(cells[j].paragraphs[0], val)
            doc.add_paragraph()
            continue

        # Bullet list
        if stripped.startswith("- "):
            while i < n and lines[i].strip().startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                _add_inline(p, lines[i].strip()[2:])
                i += 1
            continue

        # Ordered list
        if re.match(r"^\d+\.\s", stripped):
            while i < n and re.match(r"^\d+\.\s", lines[i].strip()):
                p = doc.add_paragraph(style="List Number")
                _add_inline(p, re.sub(r"^\d+\.\s", "", lines[i].strip()))
                i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        _add_inline(p, stripped)
        i += 1

    doc.save(str(docx_path))
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    convert(Path(sys.argv[1]), Path(sys.argv[2]))
